import sqlite3
import csv as csv
import itertools
import docx

#Connect to the database
conn = sqlite3.connect('GradDatabase.db')
c  = conn.cursor()


def HonorScholarList(filepath):
    c.execute("""
        SELECT Fname, MName, Lname FROM tblStudents
        WHERE HonorsScholar = 1
        ORDER BY LOWER(Lname) ASC""")
    document = docx.Document("ReportTemplates/HonorsScholar.docx")
    for item in c.fetchall():
        if item[1] is not None:
            document.add_paragraph(item[0] + " " + item[1] + " " + item[2], style = "Normal")
        else:
            document.add_paragraph(item[0] + " " + item[2], style = "Normal")

    document.save(filepath)


def boyGirlWalkingOrder(filepath):
    #SQL TO GET THE NAMES
    c.execute("""
        SELECT Fname, MName, Lname FROM tblStudents
        WHERE gender = 1
        AND Walking = 1
        AND Graduating = 1
        ORDER BY LOWER(Lname) ASC""")
    girls = c.fetchall()

    c.execute("""
        SELECT Fname, MName, Lname FROM tblStudents
        WHERE gender = 0
        AND Walking = 1
        AND Graduating = 1
        ORDER BY LOWER(Lname) ASC""")
    boys = c.fetchall()

    document = docx.Document("ReportTemplates/BoyGirlWalkingSmallFont.docx")

    for boy, girl in itertools.zip_longest( boys, girls):
        if boy is not None:
            if boy[1] is not None:
                document.add_paragraph(boy[0] + " " + boy[1] + " " + boy[2], style = "Normal")
            else:
                document.add_paragraph(boy[0] + " " + boy[2], style = "Normal")
        if girl is not None:
            if girl[1] is not None:
                document.add_paragraph(girl[0] + " " + girl[1] + " " + girl[2], style = "Normal")
            else:
                document.add_paragraph(girl[0] + " " + girl[2], style = "Normal")
    document.save(filepath)


def HonorsDiplomaList(filepath):
    c.execute("""
        SELECT Fname, MName, Lname FROM tblStudents
        WHERE HonorsDiploma = 1
        ORDER BY LOWER(Lname) ASC""")
    document = docx.Document("ReportTemplates/HonorDiploma.docx")

    for item in c.fetchall():
        if item[1] is not None:
            document.add_paragraph(item[0] + " " + item[1] + " " + item[2], style = "Normal")
        else:
            document.add_paragraph(item[0] + " " + item[2], style = "Normal")
    document.save(filepath)

def Graduating(filepath):
    c.execute("""
        SELECT Fname, MName, Lname FROM tblStudents
        WHERE Graduating = 1
        ORDER BY LOWER(Lname) ASC""")
    document = docx.Document("ReportTemplates/GraduatingAlphabetical.docx")

    for item in c.fetchall():
        if item[1] is not None:
            document.add_paragraph(item[0] + " " + item[1] + " " + item[2], style = "Normal")
        else:
            document.add_paragraph(item[0] + " " + item[2], style = "Normal")
    document.save(filepath)
